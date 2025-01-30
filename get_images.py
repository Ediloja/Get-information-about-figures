import re
import html
import requests

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.oxml import OxmlElement

from canvasapi import Canvas

from bs4 import BeautifulSoup

API_V1 = "https://utpl.instructure.com/api/v1"
API_URL = "https://utpl.instructure.com"
API_KEY = ""

HEADERS = {
    "Authorization": f"Bearer {API_KEY}",
    "Content-Type": "application/json"  # Opcional, dependiendo de la API
}

def get_connection():
    """
    Establece la conexión con la plataforma Canvas

    Parámetros:
    API_URL (str): es el dominio de Canvas
    API_KEY (str): es el token con el cual se va a trabajar

    Retorna:
    canvas (obj): Instancia a Canvas para la conexión respectiva
    """
    canvas = Canvas(API_URL, API_KEY)
   
    return canvas

def get_number(text):
    """
    Obtene un número de una cadena de  texto

    Parámetros:
    text (str): corresponde a la URL del curso

    Retorna:
    int: 
        - id_curso -> si encontró un número
        - 0 -> si no encontro números
    """

    expression = re.findall(r"\d+", text)

    if len(expression) != 0:
        return expression[0]
    else:
        return 0
    
def get_url_pages(course):
    """
    Obtiene las URLs de las páginas ordenadas

    Parámetros:
    course (obj): curso que se desea obtener la información

    Retorna:
    (list): lista de IDs de las páginas
    """

    modules = course.get_modules()

    list_pages = []

    for m in modules:
        items = m.get_module_items()

        for i in items:
            if i.type == "Page" and 'semana' in i.title.lower():
                list_pages.append(i.page_url)
            elif i.type == "Page" and 'week' in i.title.lower():
                list_pages.append(i.page_url)
    return list_pages

def get_latest_numbers(string):
    """
    Obtiene el último número de una cadena de texto

    Parámetros:
    string (str): cadena de texto que se desea extraer el último número

    Retorna:
    (int): último digito de la cadena
    (none): si no se encontró un número al final de la cadena
    """
    
    # Busca los números al final de la cadena
    match = re.search(r'\d+$', string)  
    # Devuelve los números o None si no hay
    return match.group() if match else None  

def delete_tags(html):
    """
    Elimina etiquetas del HTML "link" y "script"

    Parámetros:
    html (str): html que se desea limpiar

    Retorna:
    (str): html limpio
    """
    soup = BeautifulSoup(html, 'html.parser')

    # Se elimina las etiquetas link y script
    for tag in soup(["link", "script"]):
        tag.decompose()

    html = str(soup)
    
    return html

def identify_class(course, html):
    """
    Identifica si existe botón de continuar y se agrega el html de las páginas adicionales

    Parámetros:
    course (obj): curso que se debe analizar
    html (str): html donde se debe identificar las clases

    Retorna:
    (str): html con la información adicional (si fuera el caso)
    """

    soup = BeautifulSoup(html, 'html.parser')

    # btn_continue = soup.find_all('a', class_='boton-mas')

    # Se obtiene todas las etiquetas <a> que cumplan las siguientes condiciones: tiene la clase "boton-mas" o tiene el contenido "Leer más" o "Leer mas"
    btn_continue = [
        link for link in soup.find_all("a")
        if link.get("data-api-returntype", "").lower() == "page" 
        and "semana" not in link.text.strip().lower()
        and "week" not in link.text.strip().lower() # Excluir si contiene "semana"
    ]

    for b in btn_continue:
        url = b['data-api-endpoint']
        response = requests.get(url, headers=HEADERS)

        data = response.json()
        url_page = data.get('url')
        page = course.get_page(url_page)

        # Se incluye el html de la página externa dentro de la página
        tag_p = b.parent
        # Se limpia el html de la página externa
        html_cleaning = delete_tags(page.body)
        soup_external = BeautifulSoup(html_cleaning, 'html.parser')   

        tag_p.replace_with(soup_external)

    return str(soup)

def download_images(course, path):
    """
    Descarga las imágenes de la plataforma Canvas

    Parámetros:
    course (obj): curso que se debe analizar
    path (str): corresponde al elemento "data-api-endpoint" de la imagen

    Retorna:
    (str): nombre de la imagen. Ejemplo "figura.jpg"
    """
    id_file = get_latest_numbers(path)
    file = course.get_file(id_file)
    response = requests.get(file.url, headers=HEADERS)

    # Verificar si la descarga fue exitosa
    if response.status_code == 200:
        # Guardar la imagen localmente
        filename = file.filename
        with open(filename, "wb") as img_file:
            img_file.write(response.content)
        # Devuelve nombre de la imagen
        return filename
    else:
        print(f"Error al descargar la imagen. Código de estado: {response.status_code}")

def get_images(course, html):
    """
    Identifica nombre de las figuras, elemento, nota y texto alternativo

    Parámetros:
    course (obj): curso que se desea analizar
    html (str): html donde se debe identificar las imágenes

    Retorna:
    (str): html con la información en html con la información respectiva
    """

    soup_result = BeautifulSoup("", 'html.parser')

    soup = BeautifulSoup(html, 'html.parser')
    images = soup.find_all('img')

    for img in images:
        url = img.get('data-api-endpoint')

        # Si existe la propiedad data-api-endpoint
        if url:
            response = requests.get(url, headers=HEADERS)
            data = response.json()
            folder_id = data.get('folder_id')
            folder = course.get_folder(folder_id)

            # Si la imagen pertenece a la carpeta de "Imágenes"
            if folder.name in ["Imagenes", "Imágenes"]:
                parent_img = img.find_parent() # Se accede a la etiqueta contenedora de la imagen <p>
                title = parent_img.find_previous_sibling() # Se obtiene el título de la imagen suponiendo que es el hermano anterior

                if title:
                    # Comprobamos que no sean enlaces de las semanas
                    title_text = title.get_text().lower()
                    if any(word in title_text for word in ['figura', 'figure']):
                    # if 'figura' in title.get_text().lower() or 'figure' in title.get_text().lower():
                        soup_result.append(title)

                # Imagen
                src = img['data-api-endpoint']
                path_img = download_images(course, src)
                img_tag = soup_result.new_tag("img", src=path_img, width="300")
                img_tag_p = soup_result.new_tag("p")
                img_tag_p.append(img_tag)
                soup_result.append(img_tag_p)

                # Nota
                pre = parent_img.find_next_sibling()

                if pre and pre.name == "pre":
                    # Elimina estilos de la etiqueta pre
                    del pre['style']
                    # Reemplazamos la etiqueta por <p>
                    str_pre = str(pre).replace("<pre>", "<p>")
                    str_pre = str_pre.replace("</pre>", "</p>")
                    soup_result.append(str_pre)

                # Texto alternativo
                if 'alt' in img.attrs:
                    alt = img['alt']
                    alt_tag = soup_result.new_tag("p")
                    alt_tag.string = "Texto alternativo: " + alt
                    soup_result.append(alt_tag)
    
    return str(soup_result)

def agregar_hipervinculo(parrafo, texto, url, document):
    """
    Agrega un hipervínculo a un párrafo en un documento Word.
    
    :param parrafo: Párrafo donde se insertará el enlace.
    :param texto: Texto visible del enlace.
    :param url: URL del hipervínculo.
    :param document: Documento Word donde se inserta la relación del enlace.
    """
    # Crear una relación de hipervínculo en el documento
    r_id = document.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Crear el elemento <w:hyperlink>
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)  # Asignar ID de la relación

    # Crear el "run" (<w:r>) que contendrá el texto del enlace
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Aplicar color azul
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')

    # Aplicar subrayado
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')

    rPr.append(color)
    rPr.append(underline)
    run.append(rPr)

    # Agregar el texto al "run"
    text = OxmlElement('w:t')
    text.text = texto
    run.append(text)

    # Agregar el "run" al hipervínculo y el hipervínculo al párrafo
    hyperlink.append(run)
    parrafo._element.append(hyperlink)

def html_to_word(output_filename, html_content):
    """
    Algoritmo que convierte etiquetas html a word

    Parámetros:
    output_filename (str): ruta, nombre y extensión del archivo (ej. 'archivo.docx')
    html_content (str): html que se desea convertir a word

    Retorna:
    (file): archivo guardado
    """
    # Parsear el HTML con BeautifulSoup
    soup = BeautifulSoup(html_content, "html.parser")
    
    # Crear un documento de Word
    doc = Document()
    
    # Función para procesar elementos dentro de un párrafo
    def process_element(el, paragraph):
        if isinstance(el, str):
            paragraph.add_run(el)
        elif el.name == "strong":
            paragraph.add_run(el.get_text()).bold = True
        elif el.name == "em":
            paragraph.add_run(el.get_text()).italic = True
        elif el.name == "a":
            agregar_hipervinculo(paragraph, el.text, el.get("href"), doc)
        elif el.name == "span":
            paragraph.add_run(el.get_text())

    # Procesar cada párrafo <p> en el HTML
    for p in soup.find_all("p"):
        # Si el párrafo contiene una imagen
        if p.find("img"):
            img_tag = p.find("img")
            img_src = img_tag.get("src", "")
            try:
                doc.add_picture(img_src, width=Inches(4))  # Ajusta el ancho de la imagen
            except:
                doc.add_paragraph(f"[Imagen no encontrada: {img_src}]")  # Texto si la imagen no está disponible
        else:
            paragraph = doc.add_paragraph()
            for element in p.contents:
                process_element(element, paragraph)

    # Guardar el documento
    doc.save(output_filename)
    print(f"Documento guardado como: {output_filename}")

def write_file(filename, html):
    """
    Escribe el archivo

    Parámetros:
    filename (str): ruta, nombre y extensión del archivo (ej. 'archivo.html')
    html (str): html que se desea escribir en el archivo

    Retorna:
    (file): archivo guardado
    """
    with open(filename, "w", encoding="utf-8") as file:
        file.write(html)

    print(f"Documento guardado como: {filename}")

def decoding_html(html_text):
    """
    Limpia el HTML decodificando caracteres especiales.

    Parámetros:
    html (str): html que se desea limpiar

    Retorna:
    (str): html decodificado
    """
    return html.unescape(html_text)


def replace_br(html):
    """
    Reemplaza los saltos de línea por un punto y un espacio (título de la imagen)

    Parámetros:
    html (str): html que se desea reemplazar

    Retorna:
    (str): html sin etiquetas <br>
    """
    soup = BeautifulSoup(html, "html.parser")

    for br in soup.find_all("br"):
        br.replace_with(" ")
    
    return str(soup)

def main():
    """
        Función principal
    """

    # Se instancia la conexión a Canvas
    canvas = get_connection()

    i = 0  # Contador para el número de cursos

    # Se abre el archivo donde se encuentra el listado de cursos
    with open('courses.txt', encoding='utf8') as f:
        for line in f:
            i += 1
            url = line.strip()
            courseId = get_number(url)
            course = canvas.get_course(courseId)

            print("\n%s) %s\n" % (i, course.name))

            html_course = ""

            # Se obtiene el listado de páginas a partir de "Módulos"
            list_pages = get_url_pages(course=course)

            for p in list_pages:
                page = course.get_page(p) # Se instancia a la página
                html = delete_tags(page.body) # Se eliminan etiquetas basura
                html = identify_class(course, html) # Se suma el html de una página en particular a la página de semana
                html_course += html # se almacena el html de todas las páginas
            
            html_img = get_images(course, html_course) # Se crea un html solo: titulo, img, nota y alt
            html_dec = decoding_html(html_img) # se decodifica el html en las notas <pre>
            cleaned_html = replace_br(html_dec) # se reemplaa el br por ". "

            # Se escriben los archivos word y html
            html_to_word(f'{course.sis_course_id}.docx', cleaned_html)

if __name__ == "__main__":
    main()